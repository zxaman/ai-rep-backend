"""
Resume Parser Service
Extracts raw text from PDF and DOCX files
Phase 2: Text Extraction Layer
"""

import io
from typing import Union
import pdfplumber
from docx import Document


class ResumeParserService:
    """
    Service for extracting raw text from resume files
    Supports: PDF, DOCX
    """
    
    def __init__(self):
        """Initialize parser service"""
        self.supported_extensions = ['.pdf', '.docx', '.doc']
    
    def parse(self, file_content: bytes, filename: str) -> str:
        """
        Extract raw text from resume file
        
        Args:
            file_content: Binary file content
            filename: Original filename with extension
        
        Returns:
            Raw extracted text as string
        
        Raises:
            ValueError: If file type is unsupported
            Exception: If text extraction fails
        """
        file_ext = self._get_file_extension(filename)
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_content)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def _get_file_extension(self, filename: str) -> str:
        """Extract file extension from filename"""
        return '.' + filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    
    def _parse_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF using pdfplumber
        
        Args:
            file_content: PDF file bytes
        
        Returns:
            Extracted text from all pages
        """
        try:
            text_chunks = []
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_chunks.append(page_text)
            
            raw_text = '\n\n'.join(text_chunks)
            return raw_text.strip()
        
        except Exception as e:
            raise Exception(f"PDF parsing error: {str(e)}")
    
    def _parse_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX using python-docx
        
        Args:
            file_content: DOCX file bytes
        
        Returns:
            Extracted text from all paragraphs
        """
        try:
            doc = Document(io.BytesIO(file_content))
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_texts.append(row_text)
            
            # Combine paragraphs and tables
            all_text = paragraphs + table_texts
            raw_text = '\n\n'.join(all_text)
            
            return raw_text.strip()
        
        except Exception as e:
            raise Exception(f"DOCX parsing error: {str(e)}")
    
    def validate_file(self, filename: str) -> bool:
        """
        Check if file extension is supported
        
        Args:
            filename: Name of file to validate
        
        Returns:
            True if supported, False otherwise
        """
        file_ext = self._get_file_extension(filename)
        return file_ext in self.supported_extensions
